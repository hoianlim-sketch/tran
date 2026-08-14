import io

from docx import Document

from src.brief import parse_brief
from src.extract import load_file, load_pasted
from src.glossary import apply_glossary, parse_glossary
from src.preserve import translate_docx
from src.utils import batch_jobs, unpack_marked_units


def test_glossary_and_brief() -> None:
    pairs = parse_glossary("API = X\nboard -> board_ko\n# skip\n")
    assert pairs == [("API", "X"), ("board", "board_ko")]
    assert "X" in apply_glossary("Use the API today", pairs)

    brief = parse_brief(
        '{"tldr":"one","key_points":["a"],'
        '"outline":[{"title":"Intro","summary":"s"}],'
        '"suggested_questions":["q"]}'
    )
    assert brief.tldr == "one"
    assert brief.outline[0].title == "Intro"
    assert brief.suggested_questions == ["q"]

    spoken = parse_brief(
        '{"spoken_lines":["하나.","둘.","셋."],"tldr":"전체"}'
    )
    assert spoken.spoken_lines == ["하나.", "둘.", "셋."]
    from src.tts import spoken_script

    assert spoken_script(spoken.spoken_lines) == "하나. 둘. 셋."


def test_marked_units() -> None:
    parsed = unpack_marked_units("<<<0>>>\n안녕\n<<<2>>>\n세계")
    assert parsed[0] == "안녕"
    assert parsed[2] == "세계"
    batches = batch_jobs([(0, "a" * 10), (1, "b" * 10)], max_chars=20)
    assert len(batches) == 2


def test_load_and_preserve_docx() -> None:
    pasted = load_pasted("Hello world")
    assert pasted.method == "paste"

    source = Document()
    source.add_paragraph("Hello")
    source.add_paragraph("World")
    buffer = io.BytesIO()
    source.save(buffer)
    data = buffer.getvalue()

    loaded = load_file("sample.docx", data)
    assert "Hello" in loaded.text
    assert loaded.method == "docx"

    translated = translate_docx(data, lambda units: [item.upper() for item in units])
    result = Document(io.BytesIO(translated))
    texts = [p.text for p in result.paragraphs if p.text.strip()]
    assert texts == ["HELLO", "WORLD"]

    import pymupdf

    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Trade policy overview")
    pdf_bytes = pdf.tobytes()
    pdf.close()
    loaded_pdf = load_file("note.pdf", pdf_bytes)
    assert "Trade" in loaded_pdf.text


if __name__ == "__main__":
    test_glossary_and_brief()
    test_marked_units()
    test_load_and_preserve_docx()
    print("OK")
