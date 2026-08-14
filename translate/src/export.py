"""번역·브리핑 결과를 TXT, Markdown, DOCX로 내보낸다."""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor

from src.brief import Brief


def build_markdown(
    source_name: str,
    language: str,
    original: str,
    translation: str,
    summary: str,
    tone: str = "",
    domain: str = "",
) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    parts = [
        "# 문서 번역·브리핑 결과",
        "",
        f"- 원본: {source_name}",
        f"- 감지 언어: {language}",
        f"- 생성 시각: {now}",
    ]
    if tone:
        parts.append(f"- 문체: {tone}")
    if domain:
        parts.append(f"- 분야: {domain}")
    parts.append("")
    if summary:
        parts += ["## 한국어 브리핑", "", summary.strip(), ""]
    if translation:
        parts += ["## 한국어 번역", "", translation.strip(), ""]
    if original:
        parts += ["## 원문", "", original.strip(), ""]
    return "\n".join(parts).strip() + "\n"


def build_txt(translation: str, summary: str) -> str:
    parts: list[str] = []
    if summary:
        parts += ["[한국어 브리핑]", summary.strip(), ""]
    if translation:
        parts += ["[한국어 번역]", translation.strip()]
    return "\n".join(parts).strip() + "\n"


def build_docx(
    source_name: str,
    language: str,
    original: str,
    translation: str,
    summary: str,
    tone: str = "",
    domain: str = "",
) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    title = document.add_heading("문서 번역·브리핑 결과", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x5E)

    meta = document.add_paragraph()
    meta.add_run(f"원본: {source_name}\n").italic = True
    meta.add_run(f"감지 언어: {language}\n").italic = True
    if tone:
        meta.add_run(f"문체: {tone}\n").italic = True
    if domain:
        meta.add_run(f"분야: {domain}\n").italic = True
    meta.add_run(f"생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

    if summary:
        document.add_heading("한국어 브리핑", level=1)
        _add_body(document, summary)
    if translation:
        document.add_heading("한국어 번역", level=1)
        _add_body(document, translation)
    if original:
        document.add_heading("원문", level=1)
        _add_body(document, original)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def brief_markdown(brief: Brief | None) -> str:
    return brief.to_markdown() if brief else ""


def _add_body(document: Document, text: str) -> None:
    for block in text.strip().split("\n"):
        document.add_paragraph(block if block.strip() else "")
